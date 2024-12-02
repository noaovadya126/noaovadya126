import tkinter as tk
from tkinter import filedialog, messagebox
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import win32com.client as win32
from datetime import datetime
import re
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx import Document
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
def select_input_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.doc;*.docx")])
    input_file.set(file_path)
    if file_path:
        output_file_name = os.path.splitext(os.path.basename(file_path))[0] + "_final.docx"
        output_file.set(output_file_name)

def select_output_file():
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word files", "*.docx")],
        initialfile=output_file.get()
    )
    output_file.set(file_path)

def convert_doc_to_docx(doc_path):
    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word files", "*.docx")],
        initialfile=output_file.get()
    )
    output_file.set(file_path)


def convert_doc_to_docx(doc_path):
    try:
        # Your existing process_file code
        print("Starting file processing...")
        # Add print statements at each major step
    except Exception as e:
        print(f"Detailed error: {e}")
        print(f"Error type: {type(e)}")
        import traceback
        traceback.print_exc()  # This will print the full stack trace
    try:
        # Ensure the file path is in a proper format and check if it exists
        doc_path = os.path.abspath(doc_path)
        if not os.path.isfile(doc_path):
            raise FileNotFoundError(f"File not found: {doc_path}")

        output_path = doc_path + "x"

        # Initialize Word application
        word = win32.Dispatch("Word.Application")
        word.Visible = False

        # Try opening the document
        doc = word.Documents.Open(doc_path)
        doc.SaveAs(output_path, FileFormat=16)
        doc.Close()
        word.Quit()

        return output_path
    except Exception as e:
        # Display error message if the file cannot be opened or converted
        messagebox.showerror("Error", f"Failed to convert the .doc file to .docx.\n{e}")
        return None

def format_title_and_expand_table(doc):
    for para in doc.paragraphs:
        if "[CVR-" in para.text:
            for run in para.runs:
                run.font.size = Pt(20)
                run.font.bold = True

    for table in doc.tables:
        for row in table.rows:
            row.height = Pt(20)
            for cell in row.cells:
                cell.width = Inches(2)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)

def add_centered_header_icon(doc, icon_path):
    for section in doc.sections:
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.clear()
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = header_paragraph.add_run()
        run.add_picture(icon_path, width=Inches(1.5))



def remove_comments_and_unwanted_fields(doc):
    field_names_to_remove = [
        "Labels:", "Remaining Estimate:", "Time Spent:", "Original estimate:",
        "Assignee:", "Votes:", "Components:", "Affects versions:", "Fix versions:",
        "Request Type:", "Comments", "Comment by", "A new Customer Visit Report",
        "Installation/Maintenance approved by:", "Generated at","RDM activity description",
        "Incidentals + monetary value", "Total overnights in field", "Attachments:",
        "Type:", "Priority:", "Project:", "Any anomalies or issues to bring back to the Account Manager to discuss with the customer?:",
    ]
    field_names_approved_by = "Report approved by Neteera Head of Operations:"

    for para in doc.paragraphs:
        if any(field_name in para.text for field_name in field_names_to_remove):
            p = para._element
            p.getparent().remove(p)
            p._element = None

    for table in doc.tables:
        rows_to_remove = []
        remove_mode = False


        for i, row in enumerate(table.rows):
            row_text = " ".join(cell.text for cell in row.cells)

            if any(field_name in row_text for field_name in field_names_to_remove):
                rows_to_remove.append(row)
                remove_mode = True

            elif remove_mode:
                if any(keyword in row_text for keyword in field_names_to_remove) or not row_text.strip():
                    rows_to_remove.append(row)
                else:
                    remove_mode = False

        for row in rows_to_remove:
            row._element.getparent().remove(row._element)

def process_document_for_customer(doc):
        # Call the existing function to remove unwanted fields
        remove_comments_and_unwanted_fields(doc)

        # Additional processing: Remove `[CVR-...]` prefixes and "Due" date
        for para in doc.paragraphs:
            # Remove [CVR-...] prefix and keep the title
            if para.text.startswith("[CVR-"):
                title_start = para.text.find("]") + 1
                para.text = para.text[title_start:].strip()

            # Remove paragraphs containing "Due:" specifically
            if "Due:" in para.text:
                p = para._element
                p.getparent().remove(p)
                p._element = None

        # Remove rows in tables containing "Due:"
        for table in doc.tables:
            rows_to_remove = []

            for row in table.rows:
                row_text = " ".join(cell.text for cell in row.cells)
                if "Due:" in row_text:
                    rows_to_remove.append(row)

            # Remove marked rows
            for row in rows_to_remove:
                row._element.getparent().remove(row._element)


def modify_docx_file(file_path):
    doc = Document(file_path)
    RABNHOO = None
    last_table = None

    for table in doc.tables:  # Iterate through all tables in the document
        for row in table.rows:  # Access each row in the table
            cells = row.cells  # Get all cells in the row
            for i, cell in enumerate(cells):  # Enumerate to get the index
                if "Report approved by Neteera Head of Operations" in cell.text:
                    # Ensure there's a next cell to access
                    if i + 1 < len(cells):
                        RABNHOO = str(cells[i + 1].text.strip())  # Return the text in the adjacent cell
                    else:
                        print("Adjacent cell not found")
                    last_table = table  # Keep a reference to the table containing this text

    if last_table:
        # Add "Comments: NONE" row
        comments_row = last_table.add_row()
        comments_row.cells[0].text = "Comments:"
        comments_row.cells[1].text = "None"

        # Copy style from the first row for consistency
        for i, cell in enumerate(comments_row.cells):
            template_cell = last_table.rows[0].cells[i]
            if template_cell._element.tcPr is not None:  # Ensure tcPr exists
                cell_properties = template_cell._element.tcPr.xml  # Get the XML for tcPr
                cell._element.get_or_add_tcPr().append(
                    parse_xml(cell_properties)  # Directly append the valid XML
                )
            cell.paragraphs[0].style = template_cell.paragraphs[0].style

        # Add "Report approved..." row
        approval_row = last_table.add_row()
        approval_row.cells[0].text = "Report approved by Neteera Head of Operations:"
        approval_row.cells[1].text = RABNHOO or "Unknown"

        # Copy style from the first row for consistency
        for i, cell in enumerate(approval_row.cells):
            template_cell = last_table.rows[0].cells[i]
            if template_cell._element.tcPr is not None:  # Ensure tcPr exists
                cell_properties = template_cell._element.tcPr.xml  # Get the XML for tcPr
                cell._element.get_or_add_tcPr().append(
                    parse_xml(cell_properties)  # Directly append the valid XML
                )
            cell.paragraphs[0].style = template_cell.paragraphs[0].style
    # **************************************************************************************************************************
    for table in doc.tables:
        # Check if the table has at least one row and one cell
        if len(table.rows) > 0 and len(table.rows[0].cells) > 0:
            # Access the first cell in the first row of the table
            first_cell = table.rows[0].cells[0]
            if len(first_cell.text) > 18:
              if "Due" in first_cell.text:
                # Remove the last 16 characters from the cell's text
                first_cell.text = first_cell.text[:-17]
            if len(first_cell.text) > 18:
                if "CVR" in first_cell.text:
                    # Remove the last 16 characters from the cell's text
                    first_cell.text = first_cell.text[8:]
    # **************************************************************************************************************************
    # Save the modified document
    output_path = file_path.replace(".docx", "_modified.docx")
    doc.save(output_path)
    print(f"File successfully modified and saved as: {output_path}")
    print(f"Report approved by Neteera Head of Operations: {RABNHOO}")


def process_file():
    input_path = input_file.get()
    output_path = output_file.get()
    icon_path = os.path.join(os.path.expanduser("~"), "Downloads", "icon_x.png")

    # Basic health checks
    if not os.path.isfile(icon_path):
        messagebox.showerror("Error", f"Icon not found at {icon_path}")
        return

    if not input_path or not output_path:
        messagebox.showerror("Error", "Please select both input and output files.")
        return

    # Convert .doc to .docx if necessary
    temporary_docx_path = None
    if input_path.endswith(".doc"):
        temporary_docx_path = convert_doc_to_docx(input_path)
        input_path = temporary_docx_path  # Use the converted .docx file for processing

    doc = Document(input_path)

    add_centered_header_icon(doc, icon_path)
    format_title_and_expand_table(doc)
    remove_comments_and_unwanted_fields(doc)

    doc.save(output_path)
    #******************************
    modify_docx_file(output_path)
    #******************************
    messagebox.showinfo("Success", f"File saved as {output_path}")

    # Remove temporary files if created
    if temporary_docx_path and os.path.isfile(temporary_docx_path):
        os.remove(temporary_docx_path)


# Set up the main GUI
root = tk.Tk()
root.title("Word File Formatter")
root.geometry("600x400")  # Bigger window

input_file = tk.StringVar()
output_file = tk.StringVar()

# UI layout with padding and organization
tk.Label(root, text="Select Input File (.doc or .docx):", font=("Arial", 12, "bold")).pack(pady=10)
tk.Entry(root, textvariable=input_file, width=50, font=("Arial", 10)).pack(padx=10, pady=5)
tk.Button(root, text="Browse", command=select_input_file, font=("Arial", 10)).pack(pady=5)

tk.Label(root, text="Save Output As:", font=("Arial", 12, "bold")).pack(pady=10)
tk.Button(root, text="Save As", command=select_output_file, font=("Arial", 10)).pack(pady=5)

tk.Button(root, text="Process File", command=process_file, font=("Arial", 12, "bold"), bg="green", fg="white").pack(pady=20)

root.mainloop()
